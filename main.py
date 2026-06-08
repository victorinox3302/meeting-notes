import io
import os
import json
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
from pathlib import Path

import httpx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, File, UploadFile, HTTPException, Request, Depends, BackgroundTasks
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.security import OAuth2PasswordBearer
from dotenv import load_dotenv
from jose import JWTError, jwt
import bcrypt as _bcrypt

load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"))

app = FastAPI()

GROQ_API_KEY  = (os.getenv("GROQ_API_KEY") or "").strip()
DATABASE_URL  = os.getenv("DATABASE_URL")
SECRET_KEY    = os.getenv("SECRET_KEY", "dev-secret-change-in-production")
SMTP_EMAIL    = os.getenv("SMTP_EMAIL", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")
ALGORITHM     = "HS256"
TOKEN_DAYS    = 30

DATA_FILE     = Path(__file__).parent / "data" / "dossiers.json"
SETTINGS_FILE = Path(__file__).parent / "data" / "settings.json"

SETTINGS_DEFAULT = {
    "profil": {"prenom": "", "nom": "", "titre": "", "cabinet": ""},
    "setup": {"context": "", "folders": [], "next_id": 1}
}

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/auth/login", auto_error=False)


# ─── Auth ─────────────────────────────────────────────────────────────────────

def hash_password(password: str) -> str:
    return _bcrypt.hashpw(password.encode(), _bcrypt.gensalt()).decode()

def verify_password(plain: str, hashed: str) -> bool:
    return _bcrypt.checkpw(plain.encode(), hashed.encode())

def create_token(user_id: int) -> str:
    expire = datetime.utcnow() + timedelta(days=TOKEN_DAYS)
    return jwt.encode({"sub": str(user_id), "exp": expire}, SECRET_KEY, algorithm=ALGORITHM)

async def get_current_user(token: str = Depends(oauth2_scheme)):
    if not DATABASE_URL:
        return {"id": 0, "email": "local@dev", "role": "admin"}
    if not token:
        raise HTTPException(401, "Non authentifié")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(401, "Token invalide")
    except JWTError:
        raise HTTPException(401, "Token invalide")
    with _get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id, email, role FROM users WHERE id = %s", [int(user_id)])
            row = cur.fetchone()
            if not row:
                raise HTTPException(401, "Utilisateur introuvable")
            return {"id": row[0], "email": row[1], "role": row[2]}

def send_welcome_email(to_email: str):
    if not SMTP_EMAIL or not SMTP_PASSWORD:
        return
    try:
        html = """
        <div style="font-family:sans-serif;max-width:480px;margin:0 auto;padding:32px 24px;background:#0f172a;color:#f1f5f9;border-radius:16px;">
          <h1 style="font-size:22px;margin-bottom:8px;">Meeting Notes</h1>
          <p style="color:#94a3b8;margin-bottom:24px;">Votre assistant de consultation</p>
          <p>Bonjour,</p>
          <p>Votre compte a bien été créé. Vous pouvez dès maintenant vous connecter et utiliser l'application.</p>
          <a href="https://meeting-notes-p2yr.onrender.com"
             style="display:inline-block;margin-top:24px;padding:14px 28px;background:#6366f1;color:#fff;border-radius:12px;text-decoration:none;font-weight:600;">
            Accéder à l'application
          </a>
          <p style="margin-top:32px;color:#64748b;font-size:13px;">
            Si vous n'êtes pas à l'origine de cette inscription, ignorez cet email.
          </p>
        </div>
        """
        msg = MIMEMultipart("alternative")
        msg["Subject"] = "Bienvenue sur Meeting Notes"
        msg["From"]    = f"Meeting Notes <{SMTP_EMAIL}>"
        msg["To"]      = to_email
        msg.attach(MIMEText(html, "html"))
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(SMTP_EMAIL, SMTP_PASSWORD)
            server.sendmail(SMTP_EMAIL, to_email, msg.as_string())
    except Exception:
        pass

async def require_admin(user=Depends(get_current_user)):
    if user["role"] != "admin":
        raise HTTPException(403, "Accès réservé aux administrateurs")
    return user


# ─── Base de données ──────────────────────────────────────────────────────────

def _get_conn():
    import psycopg2
    return psycopg2.connect(DATABASE_URL)

def _init_db():
    from psycopg2.extras import Json
    with _get_conn() as conn:
        with conn.cursor() as cur:
            # Table utilisateurs
            cur.execute("""
                CREATE TABLE IF NOT EXISTS users (
                    id            SERIAL PRIMARY KEY,
                    email         TEXT UNIQUE NOT NULL,
                    password_hash TEXT NOT NULL,
                    role          TEXT NOT NULL DEFAULT 'user',
                    created_at    TIMESTAMP DEFAULT NOW()
                )
            """)

            # Détection du schéma actuel de store
            cur.execute("""
                SELECT column_name FROM information_schema.columns
                WHERE table_name = 'store'
            """)
            columns = [r[0] for r in cur.fetchall()]

            if not columns:
                # Nouvelle installation
                cur.execute("""
                    CREATE TABLE store (
                        user_id INTEGER NOT NULL REFERENCES users(id),
                        key     TEXT NOT NULL,
                        value   JSONB NOT NULL,
                        PRIMARY KEY (user_id, key)
                    )
                """)
            elif "user_id" not in columns:
                # Ancienne installation sans user_id — migration
                cur.execute("SELECT key, value FROM store")
                old_data = {r[0]: r[1] for r in cur.fetchall()}
                cur.execute("DROP TABLE store")
                cur.execute("""
                    CREATE TABLE store (
                        user_id INTEGER NOT NULL REFERENCES users(id),
                        key     TEXT NOT NULL,
                        value   JSONB NOT NULL,
                        PRIMARY KEY (user_id, key)
                    )
                """)
                # Créer l'admin
                admin_hash = hash_password("Admin2024!")
                cur.execute("""
                    INSERT INTO users (email, password_hash, role)
                    VALUES ('vhuon75@gmail.com', %s, 'admin')
                    ON CONFLICT (email) DO NOTHING
                    RETURNING id
                """, [admin_hash])
                row = cur.fetchone()
                if not row:
                    cur.execute("SELECT id FROM users WHERE email = 'vhuon75@gmail.com'")
                    admin_id = cur.fetchone()[0]
                else:
                    admin_id = row[0]
                # Migrer les données existantes vers l'admin
                for key, value in old_data.items():
                    cur.execute("""
                        INSERT INTO store (user_id, key, value) VALUES (%s, %s, %s)
                        ON CONFLICT DO NOTHING
                    """, [admin_id, key, Json(value)])

@app.on_event("startup")
def startup():
    if DATABASE_URL:
        _init_db()

def _db_load(user_id: int, key: str):
    with _get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT value FROM store WHERE user_id = %s AND key = %s", [user_id, key])
            row = cur.fetchone()
            return row[0] if row else None

def _get_admin_id() -> int:
    with _get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id FROM users WHERE role = 'admin' LIMIT 1")
            row = cur.fetchone()
            return row[0] if row else 1

def _get_admin_settings():
    s = _db_load(_get_admin_id(), "settings")
    if s is None:
        import copy
        return copy.deepcopy(SETTINGS_DEFAULT)
    s["setup"].setdefault("folders", [])
    s["setup"].setdefault("next_id", 1)
    return s

def _db_save(user_id: int, key: str, value):
    from psycopg2.extras import Json
    with _get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO store (user_id, key, value) VALUES (%s, %s, %s)
                ON CONFLICT (user_id, key) DO UPDATE SET value = EXCLUDED.value
            """, [user_id, key, Json(value)])


# ─── Accès données ────────────────────────────────────────────────────────────

def load_data(user_id=None):
    if DATABASE_URL and user_id is not None:
        data = _db_load(user_id, "dossiers")
        if data is None:
            data = {"dossiers": [], "next_id": 1001}
            _db_save(user_id, "dossiers", data)
        return data
    if not DATA_FILE.exists():
        DATA_FILE.parent.mkdir(exist_ok=True)
        DATA_FILE.write_text(json.dumps({"dossiers": [], "next_id": 1001}, ensure_ascii=False))
    return json.loads(DATA_FILE.read_text(encoding="utf-8"))

def save_data(data, user_id=None):
    if DATABASE_URL and user_id is not None:
        _db_save(user_id, "dossiers", data)
        return
    DATA_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def load_settings(user_id=None):
    if DATABASE_URL and user_id is not None:
        s = _db_load(user_id, "settings")
        if s is None:
            import copy
            s = copy.deepcopy(SETTINGS_DEFAULT)
            _db_save(user_id, "settings", s)
    else:
        if not SETTINGS_FILE.exists():
            SETTINGS_FILE.parent.mkdir(exist_ok=True)
            SETTINGS_FILE.write_text(json.dumps(SETTINGS_DEFAULT, ensure_ascii=False))
        s = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
    s["setup"].setdefault("folders", [])
    s["setup"].setdefault("next_id", 1)
    return s

def save_settings(s, user_id=None):
    if DATABASE_URL and user_id is not None:
        _db_save(user_id, "settings", s)
        return
    SETTINGS_FILE.write_text(json.dumps(s, ensure_ascii=False, indent=2), encoding="utf-8")


# ─── Static ───────────────────────────────────────────────────────────────────

app.mount("/static", StaticFiles(directory="public"), name="static")

@app.get("/")
def index():
    return FileResponse("public/index.html")


# ─── Auth endpoints ───────────────────────────────────────────────────────────

@app.post("/auth/register")
async def register(request: Request, background_tasks: BackgroundTasks):
    body = await request.json()
    email    = body.get("email", "").strip().lower()
    password = body.get("password", "").strip()
    if not email or not password:
        raise HTTPException(400, "Email et mot de passe requis")
    if len(password) < 6:
        raise HTTPException(400, "Mot de passe trop court (6 caractères minimum)")
    if not DATABASE_URL:
        raise HTTPException(500, "Auth non disponible en mode local")
    with _get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id FROM users WHERE email = %s", [email])
            if cur.fetchone():
                raise HTTPException(400, "Email déjà utilisé")
            cur.execute(
                "INSERT INTO users (email, password_hash, role) VALUES (%s, %s, 'user') RETURNING id",
                [email, hash_password(password)]
            )
            user_id = cur.fetchone()[0]
    import copy
    _db_save(user_id, "dossiers", {"dossiers": [], "next_id": 1001})
    _db_save(user_id, "settings", copy.deepcopy(SETTINGS_DEFAULT))
    background_tasks.add_task(send_welcome_email, email)
    return {"access_token": create_token(user_id), "token_type": "bearer", "role": "user"}

@app.post("/auth/login")
async def login(request: Request):
    body = await request.json()
    email    = body.get("email", "").strip().lower()
    password = body.get("password", "").strip()
    if not DATABASE_URL:
        raise HTTPException(500, "Auth non disponible en mode local")
    with _get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id, password_hash, role FROM users WHERE email = %s", [email])
            row = cur.fetchone()
    if not row or not verify_password(password, row[1]):
        raise HTTPException(401, "Email ou mot de passe incorrect")
    return {"access_token": create_token(row[0]), "token_type": "bearer", "role": row[2]}

@app.get("/auth/me")
async def me(user=Depends(get_current_user)):
    return user


# ─── Settings ─────────────────────────────────────────────────────────────────

@app.get("/settings")
def get_settings(user=Depends(get_current_user)):
    return load_settings(user["id"])

@app.post("/settings")
async def post_settings(request: Request, user=Depends(get_current_user)):
    body = await request.json()
    s = load_settings(user["id"])
    if "profil" in body:
        s["profil"].update(body["profil"])
    if "setup" in body:
        for k, v in body["setup"].items():
            if k not in ("folders", "next_id"):
                s["setup"][k] = v
    save_settings(s, user["id"])
    return s


# ─── Setup : dossiers & items (admin uniquement) ──────────────────────────────

@app.post("/settings/folders")
async def create_folder(request: Request, user=Depends(require_admin)):
    body = await request.json()
    name = body.get("name", "").strip()
    if not name:
        raise HTTPException(400, "Nom requis")
    s = load_settings(user["id"])
    folder = {"id": str(s["setup"]["next_id"]), "name": name, "items": []}
    s["setup"]["next_id"] += 1
    s["setup"]["folders"].append(folder)
    save_settings(s, user["id"])
    return folder

@app.put("/settings/folders/{fid}")
async def update_folder(fid: str, request: Request, user=Depends(require_admin)):
    body = await request.json()
    s = load_settings(user["id"])
    folder = next((f for f in s["setup"]["folders"] if f["id"] == fid), None)
    if not folder:
        raise HTTPException(404, "Dossier introuvable")
    folder["name"] = body.get("name", folder["name"]).strip()
    save_settings(s, user["id"])
    return folder

@app.delete("/settings/folders/{fid}")
async def delete_folder(fid: str, user=Depends(require_admin)):
    s = load_settings(user["id"])
    s["setup"]["folders"] = [f for f in s["setup"]["folders"] if f["id"] != fid]
    save_settings(s, user["id"])
    return {"ok": True}

@app.post("/settings/folders/{fid}/items")
async def create_item(fid: str, request: Request, user=Depends(require_admin)):
    body = await request.json()
    name = body.get("name", "").strip()
    if not name:
        raise HTTPException(400, "Nom requis")
    s = load_settings(user["id"])
    folder = next((f for f in s["setup"]["folders"] if f["id"] == fid), None)
    if not folder:
        raise HTTPException(404, "Dossier introuvable")
    item = {"id": str(s["setup"]["next_id"]), "name": name, "detail": body.get("detail", "").strip()}
    s["setup"]["next_id"] += 1
    folder["items"].append(item)
    save_settings(s, user["id"])
    return item

@app.put("/settings/folders/{fid}/items/{iid}")
async def update_item(fid: str, iid: str, request: Request, user=Depends(require_admin)):
    body = await request.json()
    s = load_settings(user["id"])
    folder = next((f for f in s["setup"]["folders"] if f["id"] == fid), None)
    if not folder:
        raise HTTPException(404, "Dossier introuvable")
    item = next((i for i in folder["items"] if i["id"] == iid), None)
    if not item:
        raise HTTPException(404, "Item introuvable")
    if "name"   in body: item["name"]   = body["name"].strip()
    if "detail" in body: item["detail"] = body["detail"].strip()
    save_settings(s, user["id"])
    return item

@app.delete("/settings/folders/{fid}/items/{iid}")
async def delete_item(fid: str, iid: str, user=Depends(require_admin)):
    s = load_settings(user["id"])
    folder = next((f for f in s["setup"]["folders"] if f["id"] == fid), None)
    if folder:
        folder["items"] = [i for i in folder["items"] if i["id"] != iid]
        save_settings(s, user["id"])
    return {"ok": True}


# ─── Dossiers ─────────────────────────────────────────────────────────────────

@app.get("/dossiers")
def list_dossiers(user=Depends(get_current_user)):
    data = load_data(user["id"])
    return {"dossiers": data["dossiers"], "next_id": data["next_id"]}

@app.post("/dossiers")
async def create_dossier(request: Request, user=Depends(get_current_user)):
    body   = await request.json()
    nom    = body.get("nom",    "").strip().upper()
    prenom = body.get("prenom", "").strip().capitalize()
    if not nom or not prenom:
        raise HTTPException(400, "Nom et prénom requis")
    data = load_data(user["id"])
    dossier = {
        "id": data["next_id"],
        "nom": nom,
        "prenom": prenom,
        "created_at": datetime.now().isoformat(),
        "setup_items": [],
        "meetings": []
    }
    data["dossiers"].append(dossier)
    data["next_id"] += 1
    save_data(data, user["id"])
    return dossier

@app.put("/dossiers/{dossier_id}")
async def update_dossier(dossier_id: int, request: Request, user=Depends(get_current_user)):
    body   = await request.json()
    nom    = body.get("nom",    "").strip().upper()
    prenom = body.get("prenom", "").strip().capitalize()
    if not nom or not prenom:
        raise HTTPException(400, "Nom et prénom requis")
    data = load_data(user["id"])
    dossier = next((d for d in data["dossiers"] if d["id"] == dossier_id), None)
    if not dossier:
        raise HTTPException(404, "Dossier introuvable")
    dossier["nom"]    = nom
    dossier["prenom"] = prenom
    save_data(data, user["id"])
    return dossier

@app.put("/dossiers/{dossier_id}/setup_items")
async def update_dossier_setup_items(dossier_id: int, request: Request, user=Depends(get_current_user)):
    body = await request.json()
    data = load_data(user["id"])
    dossier = next((d for d in data["dossiers"] if d["id"] == dossier_id), None)
    if not dossier:
        raise HTTPException(404, "Dossier introuvable")
    dossier["setup_items"] = body.get("setup_items", [])
    save_data(data, user["id"])
    return {"ok": True}

@app.post("/dossiers/{dossier_id}/meetings")
async def add_meeting(dossier_id: int, request: Request, user=Depends(get_current_user)):
    body = await request.json()
    data = load_data(user["id"])
    dossier = next((d for d in data["dossiers"] if d["id"] == dossier_id), None)
    if not dossier:
        raise HTTPException(404, "Dossier introuvable")
    meeting = {
        "date": datetime.now().isoformat(),
        "transcription": body.get("transcription", ""),
        "compte_rendu":  body.get("compte_rendu",  "")
    }
    dossier["meetings"].append(meeting)
    save_data(data, user["id"])
    return meeting

@app.patch("/dossiers/{dossier_id}/meetings/{meeting_index}")
async def update_meeting(dossier_id: int, meeting_index: int, request: Request, user=Depends(get_current_user)):
    body = await request.json()
    data = load_data(user["id"])
    dossier = next((d for d in data["dossiers"] if d["id"] == dossier_id), None)
    if not dossier:
        raise HTTPException(404, "Dossier introuvable")
    if meeting_index < 0 or meeting_index >= len(dossier["meetings"]):
        raise HTTPException(404, "Séance introuvable")
    dossier["meetings"][meeting_index]["compte_rendu"] = body.get("compte_rendu", "")
    save_data(data, user["id"])
    return dossier["meetings"][meeting_index]

@app.delete("/dossiers/{dossier_id}")
async def delete_dossier(dossier_id: int, user=Depends(get_current_user)):
    data = load_data(user["id"])
    data["dossiers"] = [d for d in data["dossiers"] if d["id"] != dossier_id]
    save_data(data, user["id"])
    return {"ok": True}


# ─── Export Word ──────────────────────────────────────────────────────────────

@app.post("/export-docx")
async def export_docx(request: Request, user=Depends(get_current_user)):
    body     = await request.json()
    text     = body.get("text", "").strip()
    patiente = body.get("patiente", "").strip()
    numero   = body.get("numero")
    date_str = body.get("date", "").strip()
    if not text:
        raise HTTPException(400, "Texte manquant")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_line = f"Compte rendu — {patiente}" if patiente else "Compte rendu"
    if patiente and numero:
        title_line += f" (N°{numero})"
    title = doc.add_heading(title_line, level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    if date_str:
        p = doc.add_paragraph(f"Séance du {date_str}")
        p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        p.runs[0].font.size = Pt(10)
        doc.add_paragraph("")

    sections = text.split("\n")
    current_title = None
    current_lines = []

    def flush_section(stitle, slines):
        if stitle:
            doc.add_heading(stitle, level=2)
        for line in slines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("• "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

    for line in sections:
        if line.startswith("## "):
            flush_section(current_title, current_lines)
            current_title = line[3:].strip()
            current_lines = []
        else:
            current_lines.append(line)
    flush_section(current_title, current_lines)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = patiente.replace(" ", "_") if patiente else "patiente"
    date_slug  = datetime.now().strftime("%Y-%m-%d")
    filename   = f"CR_{safe_name}_{date_slug}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ─── Q&A Neurosciences ────────────────────────────────────────────────────────

@app.post("/ask")
async def ask(request: Request, user=Depends(get_current_user)):
    data     = await request.json()
    question = data.get("question", "").strip()
    context  = data.get("context",  "").strip()
    if not question:
        raise HTTPException(400, "Question vide")

    system_msg = (
        "Tu es un assistant expert en neurosciences cliniques, neuropsychologie et psychiatrie. "
        "Tu aides un praticien de santé pendant la rédaction de comptes rendus de consultations. "
        "Réponds de manière concise, précise et professionnelle en français, "
        "en t'appuyant sur les données scientifiques actuelles et les meilleures pratiques cliniques."
    )
    user_msg = question
    if context:
        user_msg = f"Contexte — compte rendu de la séance en cours :\n{context}\n\nQuestion : {question}"

    async with httpx.AsyncClient() as client:
        resp = await client.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json={
                "model": "llama-3.3-70b-versatile",
                "messages": [
                    {"role": "system", "content": system_msg},
                    {"role": "user",   "content": user_msg}
                ],
                "temperature": 0.3
            },
            timeout=30.0,
        )
    result = resp.json()
    return {"answer": result["choices"][0]["message"]["content"]}


# ─── Transcription ────────────────────────────────────────────────────────────

@app.post("/transcribe")
async def transcribe(audio: UploadFile = File(...), user=Depends(get_current_user)):
    file_bytes   = await audio.read()
    filename     = audio.filename or "audio.mp3"
    content_type = audio.content_type or "audio/mpeg"

    async with httpx.AsyncClient() as client:
        resp = await client.post(
            "https://api.groq.com/openai/v1/audio/transcriptions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
            files={"file": (filename, file_bytes, content_type)},
            data={"model": "whisper-large-v3", "language": "fr"},
            timeout=120.0,
        )

    if resp.status_code != 200:
        raise HTTPException(500, f"Erreur Groq transcription : {resp.text}")

    return {"transcription": resp.json()["text"].strip()}


# ─── Structuration ────────────────────────────────────────────────────────────

@app.post("/structure")
async def structure(request: Request, user=Depends(get_current_user)):
    data          = await request.json()
    transcription = data.get("transcription", "").strip()
    prenom        = data.get("prenom", "").strip()
    nom           = data.get("nom",    "").strip()
    dossier_id    = data.get("dossier_id")
    if not transcription:
        raise HTTPException(400, "Transcription vide")

    patient_ref = f"{prenom} {nom}".strip() if prenom else "le patient"

    # Profil du praticien connecté
    user_settings = load_settings(user["id"])
    profil        = user_settings.get("profil", {})
    praticien     = " ".join(p for p in [profil.get("titre",""), profil.get("prenom",""), profil.get("nom","")] if p).strip()
    cabinet       = profil.get("cabinet", "").strip()

    # Setup IA toujours depuis le compte admin
    admin_settings = _get_admin_settings()
    setup_ctx      = admin_settings.get("setup", {}).get("context", "").strip()

    context_block = ""
    if praticien or cabinet or setup_ctx:
        context_block  = "Contexte du praticien :\n"
        if praticien: context_block += f"- Praticien : {praticien}\n"
        if cabinet:   context_block += f"- Cabinet : {cabinet}\n"
        if setup_ctx: context_block += f"- Instructions générales : {setup_ctx}\n"
        context_block += "\n"

    items_block = ""
    if dossier_id:
        try:
            doss_data = load_data(user["id"])
            doss = next((d for d in doss_data["dossiers"] if d["id"] == int(dossier_id)), None)
            if doss and doss.get("setup_items"):
                items_map = {
                    i["id"]: (f["name"], i["name"], i.get("detail", ""))
                    for f in admin_settings.get("setup", {}).get("folders", [])
                    for i in f.get("items", [])
                }
                lines = []
                for iid in doss["setup_items"]:
                    if iid in items_map:
                        fn, iname, detail = items_map[iid]
                        line = f"- {fn} / {iname}"
                        if detail:
                            line += f" : {detail}"
                        lines.append(line)
                if lines:
                    items_block = "Informations spécifiques à " + (prenom or "le patient") + " :\n" + "\n".join(lines) + "\n\n"
        except Exception:
            pass

    prompt = f"""Tu es un assistant expert en comptes rendus de consultation médicale ou de suivi patient.

{context_block}{items_block}Le patient s'appelle {patient_ref}. Dans tout le compte rendu, réfère-toi à lui uniquement par son prénom "{prenom or patient_ref}".

Voici la transcription brute de la séance :
\"\"\"{transcription}\"\"\"

Génère un compte rendu structuré, clair et professionnel en français avec exactement ces sections :

## Résumé de la séance
(2-3 phrases résumant l'essentiel, en utilisant le prénom du patient)

## Points abordés
(liste à puces)

## Observations et décisions
(liste à puces, ou "Aucune" si applicable)

## Actions à suivre
(liste à puces)

## Prochaine séance
(liste à puces ou date/objectif si mentionné)

Sois concis, factuel et professionnel. Respecte les instructions spécifiques. Utilise le prénom {prenom or "du patient"} dès que tu parles de lui."""

    async with httpx.AsyncClient() as client:
        resp = await client.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json={"model": "llama-3.3-70b-versatile", "messages": [{"role": "user", "content": prompt}], "temperature": 0.3},
            timeout=30.0,
        )
    result = resp.json()
    return {"structured": result["choices"][0]["message"]["content"]}
